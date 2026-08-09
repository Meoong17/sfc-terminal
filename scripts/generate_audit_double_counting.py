#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate styled DOCX: SFC Terminal methodology double-counting audit."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Calibri"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
HDR_FILL = "1F4E79"
HDR_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_STYLE = "Light Grid Accent 1"
OUT = "/home/ubuntu/sfc/docs/AUDIT_DOUBLE_COUNTING_METODOLOGI.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def set_heading(level, text, color=ACCENT):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT
        run.font.color.rgb = color
    return h

def para(text, bold=False, italic=False, size=11, align=None, space_after=6, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, bold_prefix=None, size=11):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r0 = p.add_run(bold_prefix); r0.bold = True
        r0.font.name = FONT; r0.font.size = Pt(size)
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(size)
    return p

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def make_table(headers, rows, widths=None, header_fill=HDR_FILL, font_size=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = TABLE_STYLE
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h); r.bold = True
        r.font.name = FONT; r.font.size = Pt(font_size); r.font.color.rgb = HDR_TEXT
        shade_cell(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.name = FONT; r.font.size = Pt(font_size)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def spacer(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    return p

# ---- title ----
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Audit Metodologi SFC Terminal\nPotensi Double-Counting / Redundant Information")
r.font.name = FONT; r.font.size = Pt(24); r.bold = True; r.font.color.rgb = ACCENT
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Komponen GLF, SLI, HMM, XGBoost, dan Composite Stress Score")
rs.font.name = FONT; rs.font.size = Pt(14); rs.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
rm = meta.add_run("Snapshot data.json 2026-08-09  |  Sumber: kode /home/ubuntu/sfc  |  bersifat audit, bukan rekomendasi investasi")
rm.font.name = FONT; rm.font.size = Pt(9); rm.italic = True; rm.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
doc.add_page_break()

# ---- 1. ringkasan eksekutif ----
set_heading(1, "1. Ringkasan Eksekutif")
para("Audit ini memetakan jalur informasi komponen GLF, SLI, HMM, XGBoost, dan skor stress "
     "komposit (effective_sfc serta composite_confidence) untuk mengidentifikasi sinyal mentah "
     "yang dihitung lebih dari satu kali di dalam agregasi yang sama.")
para("Temuan utama: headline stress (effective_sfc) adalah produk dari rantai "
     "factors[Rt,Lt,St,Mt,Ft] -> calculate_sfc_ensemble -> sfc_pct -> +liq_mod -> DW -> EWMA. "
     "Dari kelima komponen yang diaudit, SLI dan XGBoost TIDAK masuk ke skor (display-only), "
     "sehingga aman. Namun ditemukan 3 (tiga) pola double-counting pada sinyal yang benar-benar "
     "memengaruhi skor: TGA/RRP, US M2 YoY, dan DXY.", space_after=8)
make_table(
    ["Komponen", "Masuk skor?", "Kanal", "Verdict redundansi"],
    [
        ["GLF", "Ya", "factors[\"Lt\"] += glf_adj x 5.927", "Ya — sebagian komponennya terhitung ulang di kanal lain"],
        ["SLI", "Tidak", "data.json + behavioral_divergence (display-only)", "Aman"],
        ["HMM", "Indirek", "override regime -> _REGIME_DRIVER_MULT + DW", "Sebagian — fitur m2_yoy dipakai juga di GLF/liq_mod"],
        ["XGBoost", "Tidak", "blend nonaktif (weight 0.0), display-only", "Aman"],
        ["Composite stress", "Ya", "effective_sfc + composite_confidence", "Mengekor redundansi GLF/liq_mod"],
    ],
    widths=[1.0, 0.9, 2.6, 2.1],
)

# ---- 2. jalur skor utama ----
set_heading(1, "2. Jalur Skor Utama (Headline Stress)")
bullet("factors[Rt, Lt, St, Mt, Ft] dibangun dari score_factors_from_market + berbagai adjustment "
       "(ETF, FISCAL, GSLS, GLF, REPO).", bold_prefix="Step 1 — ")
bullet("calculate_sfc_ensemble(factors) menghasilkan sfc_pct dan zone (collect.py:2415).",
       bold_prefix="Step 2 — ")
bullet("sfc_pct di-nudge QLSTM (5%) lalu ditambah liq_mod (berbasis m2_yoy) menjadi effective_sfc "
       "(collect.py:2920-2924).", bold_prefix="Step 3 — ")
bullet("effective_sfc disesuaikan Dynamic Weighting (regime) dan EWMA, lalu menentukan zone/signal; "
       "composite_confidence = macro x (1 - execution_risk) (collect.py:3493).", bold_prefix="Step 4 — ")
para("Karena effective_sfc adalah bilangan tunggal tempat semua kanal bermuara, sinyal mentah yang "
     "masuk lebih dari satu kanal akan memberi dampak berlipat secara efektif.", space_after=8)

# ---- 3. temuan double-counting ----
set_heading(1, "3. Temuan Double-Counting")

set_heading(2, "3.1 TGA & RRP — dua kali masuk faktor Lt (severity tinggi)")
para("Dua modul mengambil seri FRED yang IDENTIK dan keduanya menyuntikkan hasilnya ke faktor yang sama (Lt):")
make_table(
    ["Modul", "Seri FRED", "Bobot", "Jalur ke Lt"],
    [
        ["GLF (global_liquidity_engine.py:149-150, 359-361)", "WTREGEN (TGA), RRPONTSYD (RRP)", "TGA 10%, RRP 10%", "glf_stress -> factors[\"Lt\"] (collect.py:2487)"],
        ["FISCAL M83/M84 (fiscal_liquidity.py:79, 124)", "WTREGEN (TGA), RRPONTSYD (RRP)", "M83 55%, M84 45%", "tga_adj + rrp_adj -> factors[\"Lt\"] (collect.py:2307)"],
    ],
    widths=[2.2, 1.7, 1.1, 1.6],
)
para("Satu sinyal fiskal (level + tren TGA dan RRP) diberi bobot 2x ke faktor yang sama. "
     "Snapshot live: GLF komponen tga z=-2.0 (w=0.10) dan rrp z=+1.5 (w=0.10). "
     "Pada hari drainase fiskal, GLF dan M83/M84 sama-sama mendorong Lt turun -> tekanan berlipat.", space_after=8)

set_heading(2, "3.2 US M2 YoY (M2SL) — tiga jalur, satu input (severity sedang-tinggi)")
para("Semua jalur berakar pada satu nilai m2_yoy = get_m2_data() (M2SL, collect.py:733-746). "
     "Input tunggal yang sama dikonsumsi di tiga tempat:")
make_table(
    ["Kanal", "Lokasi", "Efek"],
    [
        ["GLF komponen \"m2\" (M2SL, w=0.15)", "global_liquidity_engine.py:356-357", "Lt -> sfc_pct"],
        ["liq_mod = (7.0 - m2_yoy) x 0.8", "collect.py:2920-2924", "effective_sfc langsung (live m2_yoy=5.53 -> +1.2pp)"],
        ["HMM fitur m2_yoy", "hmm_regime.py:46", "regime -> mult + DW -> effective_sfc/zone"],
    ],
    widths=[2.2, 1.7, 2.7],
)
para("Catatan penting: komentar collect.py:2477-2480 mengklaim \"direct m2_yoy sigmoid\" sudah "
     "dihapus dari Lt sebagai redundan (diganti GLF x5.927). Namun de-duplikasi ini TIDAK TUNTAS: "
     "liq_mod di baris 2920 masih menyuntikkan m2_yoy yang sama langsung ke effective_sfc, sementara "
     "komponen m2 di dalam GLF juga masih aktif. Dampak: pertumbuhan M2 terhitung 2-3 kali.", space_after=8)

set_heading(2, "3.3 DXY — dua faktor berbeda (severity rendah-sedang)")
make_table(
    ["Kanal", "Lokasi", "Faktor tujuan"],
    [
        ["GLF komponen \"dxy\" (w=0.13)", "global_liquidity_engine.py:362", "Lt (via GLF)"],
        ["_sigmoid_factor(dxy) langsung", "collect.py:1025-1040", "factors[\"Sc\"]"],
    ],
    widths=[2.2, 1.9, 2.5],
)
para("Dua faktor (Lt dan Sc) bergerak oleh input DXY yang sama. Snapshot live: DXY=99.66 "
     "(GLF z=+0.068, netral; Sc=-0.10).", space_after=8)

# ---- 4. yang aman ----
set_heading(1, "4. Komponen yang Sudah Benar / Tidak Redundan")
make_table(
    ["Komponen", "Status", "Keterangan"],
    [
        ["execution_risk internal", "Aman", "cascade=arah, squeeze=magnitude, funding=rate asli; sudah orthogonal (de-dup 2026-08, collect.py:3440-3480)"],
        ["XGBoost meta-ensemble", "Aman", "blend dimatikan, weight 0.0, display-only (collect.py:3072)"],
        ["SLI (M76-M80)", "Aman", "display-only: data.json + behavioral_divergence, tidak masuk skor"],
        ["M72-M75 macro liquidity", "Aman (untuk skor)", "m75 dilaporkan display-only; namun inputnya (M2SL/WALCL) tumpang tindih dgn GLF"],
        ["adv_regime (k-means+Markov)", "Aman", "display-only, tidak drive scoring (per memory: tidak andal)"],
        ["M33 GLO", "Aman", "orphaned — dihitung & dilaporkan tapi TIDAK masuk ensemble (collect.py:2556-2562)"],
    ],
    widths=[2.0, 1.4, 3.2],
)

# ---- 5. rekomendasi ----
set_heading(1, "5. Rekomendasi & Keputusan (setelah walk-forward)")
bullet("Hapus adjustment FISCAL M83/M84 ke Lt (TGA/RRP kini HANYA di GLF). "
       "VALIDATED: std-ratio 0.979 (amplitudo terjaga) & polaritas justru membaik "
       "(+6.92 salah tanda -> -2.23 benar, sig). DEPLOY 2026-08-09.",
       bold_prefix="#1 — ")
bullet("Hapus liq_mod (m2_yoy langsung) dari effective_sfc. "
       "VALIDATED: std-ratio 1.206, polaritas terjaga (-2.30 -> -2.23, keduanya sig), "
       "term yang dihapus ~murni amplifikasi. DEPLOY 2026-08-09.",
       bold_prefix="#2 — ")
bullet("DXY (GLF vs Sc): PERTAHANKAN keduanya. Uji inkremental (seeded bootstrap CI 90%) "
       "menunjukkan Sc-DXY ADITIF — memperlebar gap 30d (-2.23 -> -3.14) & rentang skor "
       "(3.8-30.6 -> 2.1-57.0). Ini lensa transmisi berbeda (struktur korelasi DXY-BTC), "
       "BUKAN double-count murni. Menghapusnya akan kehilangan sinyal.",
       bold_prefix="#3 — ")
para("Setiap fiks memenuhi protokol: std-ratio OLD/NEW, polaritas prediksi, seeded numpy "
     "bootstrap, CI 90%, dan skala bobot dipertahankan (GLF x5.927 & bobot faktor lain "
     "TIDAK diubah; hanya term duplikat yang dihapus).", space_after=8)

# ---- 6. hasil walk-forward ----
set_heading(1, "6. Hasil Walk-Forward Validation (FRED 2015-2026, n=4208)")
make_table(
    ["Fiks", "std-ratio OLD/NEW", "30d gap OLD", "30d gap NEW", "Verdict"],
    [
        ["#1 TGA/RRP", "0.979", "+6.92 (salah tanda, SIG)", "-2.23 (benar, SIG)", "DEPLOY — polaritas membaik"],
        ["#2 liq_mod", "1.206", "-2.30 (SIG)", "-2.23 (SIG)", "DEPLOY — polaritas terjaga"],
        ["#3 DXY (Sc)", "std 11.85 -> 17.66", "gap base -2.23", "gap aug -3.14", "KEEP — Sc-DXY aditif"],
    ],
    widths=[1.1, 1.2, 1.6, 1.6, 2.1],
)
para("Metode: rekonstruksi FRED long-history (DTWEXBGS proxy DXY, china excluded w=0.04, "
     "St/Rt/Ft=0), formula verbatim, quantile top-vs-bottom 20% fwd-return, bootstrap seeded "
     "(random.seed(42)), nboot=20.000, CI 90%. Detail lengkap: docs/dedup_walkforward_2026.md. "
     "Catatan: rekonstruksi parsial — nilai gap absolut tidak portabel ke live; yang valid "
     "adalah perbandingan relatif OLD vs NEW (subset sama).", space_after=8)

# ---- 7. lampiran sitasi kode ----
set_heading(1, "7. Lampiran — Rujukan Kode")
make_table(
    ["Jalur", "Lokasi"],
    [
        ["GLF -> Lt (x5.927)", "collect.py:2486-2488"],
        ["FISCAL M83/M84 -> Lt", "collect.py:2304-2307"],
        ["liq_mod (m2_yoy) -> effective_sfc", "collect.py:2920-2924"],
        ["HMM fitur m2_yoy", "hmm_regime.py:46; collect.py:2946-2952"],
        ["DXY -> Sc", "collect.py:1025-1040"],
        ["GLF komponen TGA/RRP/M2/DXY", "global_liquidity_engine.py:348-363"],
        ["composite_confidence (macro x (1-exec))", "collect.py:3493"],
        ["XGBoost blend nonaktif", "collect.py:3062-3076"],
        ["Fix #1: FISCAL M83/M84 -> Lt (dihapus, display-only)", "collect.py (sebelumnya 2303-2308)"],
        ["Fix #2: liq_mod -> effective_sfc (dihapus = 0)", "collect.py (sebelumnya 2920-2924)"],
    ],
    widths=[3.6, 3.0],
)

spacer()
p = doc.add_paragraph()
r = p.add_run("DISCLAIMER: Dokumen ini adalah hasil audit metodologi berbasis kode dan data historis, "
              "bersifat informatif/riset, bukan rekomendasi investasi.")
r.font.size = Pt(8); r.italic = True; r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.save(OUT)
print("SAVED:", OUT)
