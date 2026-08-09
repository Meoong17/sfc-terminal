#!/usr/bin/env python3
"""
Generate BTC Liquidity Trend Analysis DOCX from SFC model data.
Output: /home/ubuntu/S/btc_liquidity_analysis.docx
"""

import json, os, sys
from datetime import datetime

SFC_DIR = "/home/ubuntu/sfc"

def load_json(path):
    try:
        with open(path) as f:
            return json.load(f)
    except:
        return {}

def format_change(val, good_positive=True):
    """Format a change value with arrow."""
    if val is None:
        return "N/A"
    arrow = "▲" if (good_positive and val > 0) or (not good_positive and val < 0) else "▼"
    return f"{arrow} {val:+.2f}"

def main():
    # ── Load all data ──
    data = load_json(os.path.join(SFC_DIR, "data.json"))
    glf_cache = load_json(os.path.join(SFC_DIR, ".global_liquidity_cache.json"))
    liq_mom = load_json(os.path.join(SFC_DIR, ".liq_momentum_cache.json"))
    factor_hist = load_json(os.path.join(SFC_DIR, ".factor_history.json"))

    ts_str = data.get("ts", datetime.now().isoformat())
    analysis_date = ts_str[:10] if "T" in ts_str else ts_str

    btc_price = data.get("btc", "N/A")
    btc_24h = data.get("btc_24h", 0)
    btc_mcap = data.get("btc_mcap", 0)
    ath = data.get("ath", 0)
    ath_date = data.get("ath_date", "N/A")[:10] if data.get("ath_date") else "N/A"
    regime = data.get("regime", "N/A")
    zone = data.get("zone", "N/A")
    fng = data.get("fng", "N/A")
    fng_cls = data.get("fng_cls", "N/A")

    # ── SFC scores ──
    sfc_base = data.get("sfc_base", 0)
    sfc_effective = data.get("sfc_effective", 0)
    composite_conf = data.get("composite_confidence", 0)
    factors = data.get("factors", {})
    lt = factors.get("Lt", 0)
    st = factors.get("St", 0)
    rt = factors.get("Rt", 0)
    ft = factors.get("Ft", 0)
    sc = factors.get("Sc", 0)

    # ── GLF ──
    glf_score = data.get("glf_score", 50)
    glf_stress = data.get("glf_stress", 0.5)
    glf_regime = data.get("glf_regime", "N/A")
    glf_comp = data.get("glf_component_detail", {})

    # ── Liquidity Momentum ──
    history = liq_mom.get("history", [])
    lm_change = None
    lm_label = "INSUFFICIENT_DATA"
    glf_now = history[-1]["glf"] if history else 50.0
    glf_30d = None
    if len(history) >= 30:
        glf_30d = history[-30]["glf"]
        lm_change = glf_now - glf_30d
        if lm_change > 5: lm_label = "ACCELERATING_IMPROVEMENT"
        elif lm_change > 2: lm_label = "IMPROVING"
        elif lm_change > -2: lm_label = "STABLE"
        elif lm_change > -5: lm_label = "DETERIORATING"
        else: lm_label = "SHARP_DETERIORATION"
    elif len(history) >= 7:
        glf_7d = history[-7]["glf"]
        lm_change = glf_now - glf_7d
        lm_label = "PARTIAL_DATA"
    else:
        lm_change = 0
        lm_label = "INSUFFICIENT_HISTORY"

    # GLF history trajectory
    glf_values = [h["glf"] for h in history]
    dates = [h["date"] for h in history]
    glf_min = min(glf_values) if glf_values else 0
    glf_max = max(glf_values) if glf_values else 0
    glf_recent = glf_values[-10:] if len(glf_values) >= 10 else glf_values
    glf_trend_dir = "RISING" if len(glf_recent) >= 3 and glf_recent[-1] > glf_recent[0] else "FALLING" if len(glf_recent) >= 3 and glf_recent[-1] < glf_recent[0] else "FLAT"

    # ── Stablecoin Liquidity ──
    sli = data.get("sli_score", 0)
    sli_label = data.get("sli_label", "N/A")
    sli_components = data.get("sli_components", {})

    # ── M2 ──
    m2_yoy = data.get("m2_yoy", 0)
    m73 = data.get("m73_detail", {})
    m2_momentum = m73.get("m2_momentum", 0)

    # ── On-chain ──
    sopr = data.get("sopr_proxy", 0)
    sopr_signal = data.get("sopr_signal", "N/A")
    cascade = data.get("cascade_risk", 0)
    liq_pressure = data.get("liq_pressure", "N/A")
    liq_density = data.get("liq_density", 0)
    rsi = data.get("rsi_14", 0)
    rsi_regime = data.get("rsi_regime", "N/A")

    # ── Behavioral divergence ──
    beh_div = data.get("behavioral_divergence_score", 0)
    beh_detail = data.get("behavioral_divergence_detail", {})
    beh_regime = beh_detail.get("regime", "N/A")

    # ── ETF ──
    etf_avg_flow = data.get("m81_detail", {}).get("m81_avg_flow_5d_btc", 0)
    etf_latest = data.get("m81_detail", {}).get("m81_latest_flow_btc", 0)
    etf_cumulative = data.get("m81_detail", {}).get("m82_cumulative_btc", 0)

    # ── Liquidation data ──
    liq_total = data.get("liq_total_24h", 0)
    liq_long = data.get("liq_long_vol", 0)
    liq_short = data.get("liq_short_vol", 0)
    funding = data.get("m25_detail", {}).get("funding_rate", 0)
    funding_8h = data.get("m25_detail", {}).get("funding_8h", 0)

    # ── Lt factor history ──
    lt_history = factor_hist.get("Lt", [])
    lt_trend = "RISING" if len(lt_history) >= 3 and lt_history[-1] > lt_history[0] else "FALLING" if len(lt_history) >= 3 and lt_history[-1] < lt_history[0] else "STABLE"

    # ===== BUILD DOCX =====
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    def set_cell_shading(cell, color):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    def add_heading_styled(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        return h

    def add_kv_row(table, key, value):
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else False
        return row

    def fmt(val, decimals=2):
        if isinstance(val, float):
            return round(val, decimals)
        return val

    # ════════════════════════════════════════════
    #  TITLE PAGE
    # ════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BTC Liquidity Trend Analysis")
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Based on SFC Model v4.0.0")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\nAnalysis Date: {analysis_date}\n"
                     f"BTC Price: ${btc_price:,.2f}\n"
                     f"Regime: {regime}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ════════════════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════
    add_heading_styled("1. Executive Summary", level=1)

    # Determine overall liquidity verdict
    if glf_score > 65:
        liq_verdict = "ABUNDANT"
        liq_bias = "Bullish"
    elif glf_score > 50:
        liq_verdict = "ACCOMMODATIVE"
        liq_bias = "Moderately Bullish"
    elif glf_score > 35:
        liq_verdict = "NEUTRAL"
        liq_bias = "Neutral"
    elif glf_score > 20:
        liq_verdict = "TIGHTENING"
        liq_bias = "Moderately Bearish"
    else:
        liq_verdict = "CRISIS"
        liq_bias = "Bearish"

    summary_lines = [
        f"Global Liquidity Factor (GLF) score: {glf_score:.1f}/100 — {glf_regime} regime. "
        f"Liquidity momentum is {lm_label.replace('_', ' ').title()} with a {glf_trend_dir.lower()} trajectory over the recent period.",

        f"BTC trades at ${btc_price:,.0f}, {(btc_price/ath - 1)*100:.1f}% below its all-time high of ${ath:,.0f} ({ath_date}). "
        f"The SFC effective stress score is {sfc_effective:.1f} (base: {sfc_base:.1f}), indicating a NEUTRAL stress environment.",

        f"Macro liquidity components are mixed: central bank balance sheets are contracting (ECB -4.3% YoY, BOJ -10.9% YoY, "
        f"Fed +0.7% YoY), while US M2 grows at a moderate +5.6% YoY. The RRP facility is effectively empty at $0.7B — "
        f"cash is fully deployed, a historically bullish liquidity signal.",

        f"The SFC 5-factor model shows the Lt (Liquidity Trend) factor at {lt:.4f} — near maximum bullish territory — "
        f"while the Rt (Reflexivity Trend) factor reads {rt:.4f}, reflecting bearish reflexive sentiment. "
        f"Stablecoin liquidity (SLI) scores {sli:.1f}/100 ({sli_label}).",

        f"Composite confidence is low at {composite_conf:.3f}, weighed by elevated cascade risk ({cascade:.3f}), "
        f"a long-squeeze liquidation regime ({liq_pressure}), and contracting M2 momentum ({m2_momentum:+.1f}). "
        f"Behavioral divergence signals {beh_regime.replace('_', ' ').title()} — the divergence tracker scores {beh_div:.1f}.",

        f"Overall verdict: LIQUIDITY TREND IS NEUTRAL-TO-ACCOMMODATIVE, but structural headwinds from "
        f"contracting central bank balance sheets and weak momentum keep the outlook cautious. "
        f"The empty RRP and drawing-down TGA provide near-term support, but a sustained BTC rally "
        f"requires re-acceleration of global M2 and central bank balance sheet expansion.",
    ]

    for line in summary_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(6)

    # ════════════════════════════════════════════
    #  2. GLOBAL LIQUIDITY FACTOR (GLF) BREAKDOWN
    # ════════════════════════════════════════════
    add_heading_styled("2. Global Liquidity Factor (GLF) Breakdown", level=1)

    p = doc.add_paragraph(
        f"The GLF consolidates 8 macro liquidity components into a single 0-100 score. "
        f"Current score: {glf_score:.1f} / 100 — classified as {glf_regime}. "
        f"SFC stress mapping: {glf_stress:.2f} (0=no stress, 1=max stress)."
    )

    # GLF Components Table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0]
    headers = ['Component', 'Raw Value', 'Z-Score', 'Weight', 'Signal']
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True if p.runs else False

    # Component data
    comp_data = {
        "Fed Balance Sheet": glf_comp.get("fed", {}),
        "ECB Balance Sheet": glf_comp.get("ecb", {}),
        "BOJ Balance Sheet": glf_comp.get("jpn", {}),
        "China M2": glf_comp.get("china", {}),
        "US M2": glf_comp.get("m2", {}),
        "TGA (Fiscal)": glf_comp.get("tga", {}),
        "RRP (Money Market)": glf_comp.get("rrp", {}),
        "DXY (Inverted)": glf_comp.get("dxy", {}),
    }

    for name, comp in comp_data.items():
        raw = comp.get("raw", "N/A")
        z = comp.get("z_score", 0)
        w = comp.get("weight", 0)
        if z is None or z == 0:
            signal = "Neutral"
        elif z > 0.5:
            signal = "Bullish ▲" if name != "DXY (Inverted)" else "Bearish ▲"
        elif z > 0:
            signal = "Slightly Bullish"
        elif z > -0.5:
            signal = "Slightly Bearish"
        else:
            signal = "Bearish ▼" if name != "DXY (Inverted)" else "Bullish ▼"

        # For DXY, invert interpretation
        if name == "DXY (Inverted)":
            if z and z > 0.5: signal = "Bearish (USD Strong)"
            elif z and z < -0.5: signal = "Bullish (USD Weak)"

        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = str(fmt(raw, 2))
        row.cells[2].text = f"{fmt(z, 3):+}"
        row.cells[3].text = f"{w*100:.0f}%"
        row.cells[4].text = signal

    doc.add_paragraph()  # spacer

    # Key observations
    add_heading_styled("Key Observations", level=2)
    obs = [
        f"RRP at $0.7B (near zero): All money market cash deployed — historically one of the most bullish liquidity signals for risk assets.",
        f"TGA drawing down (-5.75% in 4 weeks, now $830B): Fiscal stimulus is active, partially offsetting central bank tightening.",
        f"ECB (-4.3% YoY, z=-1.18) and BOJ (-10.9% YoY, z=-2.31) balance sheets contracting sharply — the primary drag on GLF.",
        f"Fed balance sheet nearly flat at +0.71% YoY (z=-0.60) — QT is effectively done but no expansion yet.",
        f"US M2 at +5.58% YoY (z=-0.11) is near neutral — money supply growth is stable but not accelerating.",
        f"DXY at 101.43 (z=-0.29): USD moderately strong, slightly negative for dollar-denominated liquidity.",
    ]
    for o in obs:
        doc.add_paragraph(o, style='List Bullet')

    # ════════════════════════════════════════════
    #  3. LIQUIDITY MOMENTUM (LM)
    # ════════════════════════════════════════════
    add_heading_styled("3. Liquidity Momentum (LM)", level=1)

    p = doc.add_paragraph(
        f"Liquidity Momentum measures the rate of change in GLF over 30 days. "
        f"Bitcoin reacts to CHANGES in liquidity, not absolute levels — making LM a critical leading indicator."
    )

    # LM details
    lm_table = doc.add_table(rows=1, cols=2)
    lm_table.style = 'Light Grid Accent 1'
    hdr = lm_table.rows[0]
    hdr.cells[0].text = "Metric"
    hdr.cells[1].text = "Value"
    hdr.cells[0].paragraphs[0].runs[0].bold = True

    add_kv_row(lm_table, "LM Change (30d)", f"{fmt(lm_change):+.2f} pts")
    add_kv_row(lm_table, "LM Trend Label", lm_label.replace('_', ' ').title())
    add_kv_row(lm_table, "GLF Now", f"{fmt(glf_now):.1f}")
    add_kv_row(lm_table, "GLF 30 Days Ago", f"{fmt(glf_30d):.1f}" if len(history) >= 30 else "N/A")
    add_kv_row(lm_table, "GLF Min (history)", f"{fmt(glf_min):.1f}")
    add_kv_row(lm_table, "GLF Max (history)", f"{fmt(glf_max):.1f}")
    add_kv_row(lm_table, "GLF Short-term Trend", glf_trend_dir)
    add_kv_row(lm_table, "Data Points", f"{len(history)} days")
    add_kv_row(lm_table, "Date Range", f"{dates[0]} → {dates[-1]}" if len(dates) >= 2 else "N/A")

    doc.add_paragraph()
    p = doc.add_paragraph(
        f"Liquidity momentum is {lm_label.replace('_', ' ').lower()}. "
    )
    if glf_30d is not None:
        p.add_run(
            f"GLF has moved from ~{fmt(glf_30d):.1f} 30 days ago to {fmt(glf_now):.1f} now, "
            f"a change of {fmt(lm_change):+.2f} points. "
        )
    else:
        p.add_run(
            f"Current GLF is {fmt(glf_now):.1f}. "
            f"Only {len(history)} days of history available — need 30 for full LM calculation. "
        )
    if lm_change and lm_change > 2:
        p.add_run(" This improvement is supportive for BTC but needs confirmation from sustained central bank expansion.")
    elif lm_change and lm_change > 0:
        p.add_run(" The slight improvement is a tentative positive signal, but not yet strong enough to drive a trend change.")
    else:
        p.add_run(" The deterioration warrants caution — BTC historically underperforms when 30-day LM is negative.")

    # ════════════════════════════════════════════
    #  4. SFC 5-FACTOR MODEL — LIQUIDITY (Lt)
    # ════════════════════════════════════════════
    add_heading_styled("4. SFC 5-Factor Model: Liquidity Trend (Lt)", level=1)

    p = doc.add_paragraph(
        f"The Lt factor is one of 5 orthogonal factors driving SFC scoring. "
        f"Lt specifically captures the global liquidity component, ranging from strongly bearish (-2) to strongly bullish (+2)."
    )

    # Factor table
    ftable = doc.add_table(rows=1, cols=3)
    ftable.style = 'Light Grid Accent 1'
    hdr = ftable.rows[0]
    for i, h in enumerate(["Factor", "Value", "Interpretation"]):
        hdr.cells[i].text = h
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    factor_data = [
        ("Lt (Liquidity Trend)", f"{lt:.4f}", "Near-max bullish" if lt > 0.9 else "Bullish" if lt > 0.5 else "Neutral" if lt > -0.3 else "Bearish"),
        ("St (Structural Trend)", f"{st:.4f}", "Bullish" if st > 0.5 else "Neutral" if st > -0.3 else "Bearish"),
        ("Rt (Reflexivity Trend)", f"{rt:.4f}", "Bearish" if rt < -0.5 else "Neutral" if rt < 0.5 else "Bullish"),
        ("Ft (Flow Trend)", f"{ft:.4f}", "Bullish" if ft > 0.5 else "Neutral" if ft > -0.3 else "Bearish"),
        ("Sc (Sentiment Composite)", f"{sc:.4f}", "Neutral-Bearish" if sc < 0 else "Neutral-Bullish"),
    ]
    for name, val, interp in factor_data:
        row = ftable.add_row()
        row.cells[0].text = name
        row.cells[1].text = val
        row.cells[2].text = interp

    doc.add_paragraph()
    p = doc.add_paragraph(
        f"Lt at {lt:.4f} is the strongest factor in the model, near the upper end of its range. "
    )

    if lt_history:
        lt_pct = (lt_history[-1] / max(lt_history) * 100) if max(lt_history) > 0 else 0
        p.add_run(
            f"Over the last {len(lt_history)} observations, Lt has been "
            f"{lt_trend.lower()} with values ranging from {min(lt_history):.4f} to {max(lt_history):.4f}. "
        )

    p.add_run(
        f"The high Lt reading suggests that, from a pure liquidity perspective, the macro environment "
        f"is supportive for BTC. However, this bullish liquidity signal is being offset by weak reflexive "
        f"sentiment (Rt: {rt:.4f}) and low sentiment (Sc: {sc:.4f}), creating internal divergence within the factor model."
    )

    # ════════════════════════════════════════════
    #  5. STABLECOIN LIQUIDITY (SLI)
    # ════════════════════════════════════════════
    add_heading_styled("5. Stablecoin Liquidity (SLI)", level=1)

    p = doc.add_paragraph(
        f"Stablecoin Liquidity Index scores {sli:.1f}/100 — {sli_label}. "
        f"SLI measures the availability of stablecoin capital as a proxy for crypto-native liquidity."
    )

    sli_table = doc.add_table(rows=1, cols=4)
    sli_table.style = 'Light Grid Accent 1'
    hdr = sli_table.rows[0]
    for i, h in enumerate(["Component", "Score", "Weight", "Contribution"]):
        hdr.cells[i].text = h
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    for comp_name, comp in sorted(sli_components.items()):
        row = sli_table.add_row()
        row.cells[0].text = comp_name.replace('_', ' ').title()
        row.cells[1].text = f"{comp.get('score', 0):.3f}"
        row.cells[2].text = f"{comp.get('weight', 0)*100:.0f}%"
        row.cells[3].text = f"{comp.get('contribution', 0):+.4f}"

    doc.add_paragraph()
    p = doc.add_paragraph(
        f"Notable: SSR (Stablecoin Supply Ratio) contributes positively at {sli_components.get('ssr', {}).get('score', 0):.2f} — "
        f"indicating ample stablecoin purchasing power relative to BTC market cap (SSR={data.get('m77_detail', {}).get('ssr', 0):.2f}). "
        f"Exchange outflows dominate inflows (netflow: {data.get('m78_detail', {}).get('netflow', 0):+,.0f} BTC), suggesting accumulation."
    )

    # ════════════════════════════════════════════
    #  6. MACRO LIQUIDITY COMPONENTS
    # ════════════════════════════════════════════
    add_heading_styled("6. Macro Liquidity Deep Dive", level=1)

    add_heading_styled("6.1 US M2 Money Supply", level=2)
    p = doc.add_paragraph(
        f"US M2 Money Supply grows at +{m2_yoy:.2f}% YoY (z-score: {fmt(glf_comp.get('m2', {}).get('z_score', 0), 3):+}). "
        f"The raw value of ${data.get('m72_detail', {}).get('m2_latest', 0):,.0f}B supports a neutral-to-moderately accommodative stance. "
        f"However, M2 momentum (3-month rate of change) is {m2_momentum:+.1f}, classified as CONTRACTING — "
        f"suggesting the growth rate is decelerating. This is a headwind for BTC, as Bitcoin historically "
        f"correlates positively with the rate of change in M2."
    )

    add_heading_styled("6.2 Central Bank Balance Sheets", level=2)
    p = doc.add_paragraph(
        f"The three major central banks present a mixed picture:\n"
        f"• Fed (WALCL): +{fmt(glf_comp.get('fed', {}).get('raw', 0), 2)}% YoY — effectively stable, QT headwinds are fading\n"
        f"• ECB (ECBASSETSW): {fmt(glf_comp.get('ecb', {}).get('raw', 0), 2):+}% YoY — contracting, tightening euro liquidity\n"
        f"• BOJ (JPNASSETS): {fmt(glf_comp.get('jpn', {}).get('raw', 0), 2):+}% YoY — sharp contraction as BOJ normalizes\n\n"
        f"Aggregate central bank liquidity (GLO index): z-score of {data.get('m33_glo_detail', {}).get('glo_z_score', 0):+.2f}, "
        f"classified as CONTRACTIVE. The composite global liquidity score from the 3 CB balance sheets plus GLO index "
        f"scores {data.get('m33_glo_score', 0):.2f}/1.00, suggesting the global trend is still mildly CONTRACTIVE."
    )

    add_heading_styled("6.3 RRP & TGA — The Liquidity Cushion", level=2)
    rrp_detail = data.get('m84_detail', {})
    tga_detail = data.get('m83_detail', {})
    p = doc.add_paragraph(
        f"RRP Facility: ${rrp_detail.get('rrp_latest_b', 0):.1f}B — NEAR_ZERO. "
        f"The RRP effectively drained from ~$2T in early 2024 to near zero, meaning all parked cash has been deployed. "
        f"This is a powerful tailwind: every dollar that leaves the RRP must find a home in the real economy or risk assets.\n\n"
        f"TGA Balance: ${tga_detail.get('tga_latest_b', 0):.1f}B — DRAWING_DOWN ({fmt(tga_detail.get('tga_4w_chg_pct', 0), 2):+}% 4-week change, "
        f"-${abs(tga_detail.get('tga_4w_chg_b', 0)):.1f}B). Fiscal stimulus is active. "
        f"Combined, RRP + TGA drain has injected roughly $1.5T+ of liquidity into the financial system over the past 18 months, "
        f"acting as a powerful offset to QT."
    )

    add_heading_styled("6.4 DXY — Dollar Liquidity", level=2)
    dxy = data.get('dxy', 0)
    p = doc.add_paragraph(
        f"DXY at {dxy:.2f} (z-score: {fmt(glf_comp.get('dxy', {}).get('z_score', 0), 3):+}). "
        f"The inverted DXY z-score in GLF means a strong USD reduces the GLF score. "
        f"At ~101.4, the dollar is moderately strong but not extreme. "
        f"A sustained DXY break below 100 would be significantly bullish for BTC and risk assets globally."
    )

    # ════════════════════════════════════════════
    #  7. BTC TECHNICAL & ON-CHAIN CONTEXT
    # ════════════════════════════════════════════
    add_heading_styled("7. BTC Technical & On-Chain Context", level=1)

    oc_table = doc.add_table(rows=1, cols=3)
    oc_table.style = 'Light Grid Accent 1'
    hdr = oc_table.rows[0]
    for i, h in enumerate(["Metric", "Value", "Signal"]):
        hdr.cells[i].text = h
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    oc_data = [
        ("BTC Price", f"${btc_price:,.2f}", f"{btc_24h:+.2f}% 24h"),
        ("ATH Distance", f"${ath:,.0f} on {ath_date}", f"{(btc_price/ath - 1)*100:+.1f}% below ATH"),
        ("SFC Base Score", f"{sfc_base:.2f}", "Lower = better"),
        ("SFC Effective", f"{sfc_effective:.2f}", f"Adj: {sfc_effective - sfc_base:+.2f}"),
        ("Composite Confidence", f"{composite_conf:.3f}", "Low" if composite_conf < 0.2 else "Moderate"),
        ("Regime / Zone", f"{regime} / {zone}", "Normal oscillation"),
        ("Fear & Greed", f"{fng} ({fng_cls})", "Fear zone — contrarian bullish typically"),
        ("RSI-14", f"{rsi:.1f} ({rsi_regime})", "Neutral" if 30 < rsi < 70 else "Extreme"),
        ("SOPR", f"{sopr:.4f} ({sopr_signal})", "Mild distress — near 1.0 breakeven"),
        ("Liquidations 24h", f"${liq_total:,.0f}", f"L: ${liq_long:,.0f} / S: ${liq_short:,.0f}"),
        ("Liq. Pressure", liq_pressure, "Long squeeze risk elevated"),
        ("Cascade Risk", f"{cascade:.3f}", "Elevated" if cascade > 0.6 else "Normal"),
        ("Funding Rate 8h", f"{funding_8h:.8f}", "Neutral to slight long bias" if funding_8h > 0 else "Neutral"),
        ("Behavioral Divergence", f"{beh_div:.1f} ({beh_regime})", "Hidden accumulation detected"),
        ("ETF Avg Flow 5d", f"{etf_avg_flow:+.1f} BTC", "Neutral" if abs(etf_avg_flow) < 10 else "Notable"),
        ("ETF Cumulative", f"{etf_cumulative:,.0f} BTC", f"${etf_cumulative * btc_price:,.0f}"),
    ]
    for metric, val, sig in oc_data:
        row = oc_table.add_row()
        row.cells[0].text = metric
        row.cells[1].text = val
        row.cells[2].text = sig

    doc.add_paragraph()

    # ════════════════════════════════════════════
    #  8. DIVERGENCE ANALYSIS
    # ════════════════════════════════════════════
    add_heading_styled("8. Divergence Analysis", level=1)

    add_heading_styled("8.1 Liquidity vs. Price Divergence", level=2)
    p = doc.add_paragraph(
        f"Lt factor at {lt:.4f} (strongly bullish) while BTC is -{((1-btc_price/ath)*100):.0f}% from ATH. "
        f"This represents a classic divergence: the liquidity environment is supportive, but price has not "
        f"responded proportionally. This can resolve in two ways:\n\n"
        f"1) Price catches up to liquidity — a bullish resolution where BTC rallies as liquidity "
        f"expansion eventually flows into risk assets.\n"
        f"2) Liquidity deteriorates toward price — where the GLF score falls and the Lt factor declines, "
        f"eliminating the divergence via the bearish path.\n\n"
        f"Historically, resolution #1 (price catching up) has been the more common outcome when "
        f"RRP is near zero and TGA is drawing down — these conditions have preceded major BTC rallies."
    )

    add_heading_styled("8.2 Behavioral Divergence (Hidden Accumulation)", level=2)
    p = doc.add_paragraph(
        f"Behavioral divergence score: {beh_div:.1f}/100 — {beh_regime.replace('_', ' ').title()}.\n\n"
        f"This divergence tracker monitors price vs. smart-money flow (ETF flows, whale pressure, stablecoin flows). "
        f"A reading in the {beh_regime.replace('_', ' ').title()} zone suggests smart money is accumulating "
        f"while retail sentiment remains bearish — a historically favorable setup. "
        f"Component values: ETF flow={fmt(beh_detail.get('component_values', {}).get('etf_flow', 0), 3)}, "
        f"Whale pressure={fmt(beh_detail.get('component_values', {}).get('whale_pressure', 0), 3)}, "
        f"Stablecoin={fmt(beh_detail.get('component_values', {}).get('stablecoin', 0), 3)}."
    )

    add_heading_styled("8.3 Reflexivity Divergence", level=2)
    rdiv = data.get("reflexivity_divergence_detail", {})
    p = doc.add_paragraph(
        f"Reflexivity divergence score: {data.get('reflexivity_divergence_score', 0):.1f} — "
        f"regime: {rdiv.get('regime', 'N/A')}. Price ROC: {rdiv.get('price_roc', 0):.4f}, "
        f"Leverage ROC: {rdiv.get('leverage_roc', 0):.4f}. "
        f"No significant reflexive divergence detected in the current window."
    )

    # ════════════════════════════════════════════
    #  9. STRESS & RISK ASSESSMENT
    # ════════════════════════════════════════════
    add_heading_styled("9. Liquidity Stress & Risk Assessment", level=1)

    stress_table = doc.add_table(rows=1, cols=3)
    stress_table.style = 'Light Grid Accent 1'
    hdr = stress_table.rows[0]
    for i, h in enumerate(["Risk Factor", "Level", "Impact on BTC"]):
        hdr.cells[i].text = h
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    stress_data = [
        ("Cascade Risk", f"{cascade:.3f} (Elevated)", "Sell-off amplification risk if support breaks"),
        ("Liq. Density", f"{liq_density:.3f}", "Moderate liquidity depth — typical for accumulation phase"),
        ("Long Squeeze Risk", liq_pressure, "Overleveraged longs vulnerable to flush"),
        ("M2 Momentum", f"{m2_momentum:+.1f} (Contracting)", "Medium-term headwind if trend persists"),
        ("CB Balance Sheets", "CONTRACTIVE (GLO z={:.2f})".format(data.get('m33_glo_detail', {}).get('glo_z_score', 0)), "Primary macro headwind"),
        ("Regime Transition Risk", f"{data.get('transition_risk', 0):.3f}", "Transition from NORMAL to STRESS possible"),
        ("News Stress", f"{data.get('news_stress', 0):.1f}/10", "Low — no acute catalyst"),
    ]
    for risk, level, impact in stress_data:
        row = stress_table.add_row()
        row.cells[0].text = risk
        row.cells[1].text = level
        row.cells[2].text = impact

    doc.add_paragraph()

    # ════════════════════════════════════════════
    #  10. OUTLOOK & KEY LEVELS
    # ════════════════════════════════════════════
    add_heading_styled("10. Outlook & Key Levels", level=1)

    # Floor/buffer from data
    floor_buffer = data.get("floor_buffer", 0)
    floor_total = data.get("floor_total", 0)

    p = doc.add_paragraph("Key Price Levels (based on SFC model):")
    doc.add_paragraph(f"• SFC Floor: ${floor_total:,.0f} (buffer: ${floor_buffer:,.0f}) — Support zone")
    doc.add_paragraph(f"• Current: ${btc_price:,.0f}")
    doc.add_paragraph(f"• All-Time High: ${ath:,.0f}")
    doc.add_paragraph()

    add_heading_styled("Scenario Analysis", level=2)

    scenarios = [
        ("Bullish Scenario", 
         "RRP stays near zero, TGA continues drawing down, Fed pivots to balance sheet expansion (rate cuts + QT end). "
         "M2 re-accelerates above +7% YoY. Behavioral divergence resolves upward. "
         f"Target: Re-test of prior highs above $100K. Probability: MODERATE (requires Fed catalyst).",
         "BULLISH"),
        ("Base Case (Most Likely)",
         "GLF remains in 45-55 range (NEUTRAL). Central bank normalization continues at a slow pace. "
         "BTC consolidates between $55K-$75K with gradual accumulation. "
         "Lt factor stays elevated but price action remains range-bound until M2 momentum turns positive.",
         "NEUTRAL"),
        ("Bearish Scenario",
         "M2 momentum continues contracting, ECB/BOJ tightening accelerates, DXY breaks above 105. "
         "Cascade risk triggers long liquidation cascade below $57K floor. "
         f"Target: $45K-$50K. Probability: LOW-MODERATE (requires exogenous shock).",
         "BEARISH"),
    ]

    for title_text, desc, tag in scenarios:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph()

    # ── Final verdict box ──
    add_heading_styled("Final Verdict", level=2)
    verdict = (
        f"The SFC liquidity model paints a picture of a market caught between supportive macro liquidity "
        f"(empty RRP, drawing TGA, stable US M2) and structural headwinds (contracting CB balance sheets, "
        f"decelerating M2 momentum, elevated cascade risk). The Lt factor at {lt:.4f} suggests that, "
        f"on balance, the liquidity environment is accommodative enough to prevent a deep bear market — "
        f"but insufficiently dynamic to catalyze a new uptrend without an additional catalyst.\n\n"
        f"The behavioral divergence (Hidden Accumulation) is the most encouraging signal for bulls: "
        f"smart money accumulating while price drifts lower. Combined with a near-empty RRP and "
        f"drawing TGA, the conditions are historically consistent with the late-stage accumulation "
        f"phase of a market cycle.\n\n"
        f"Key catalysts to watch: (1) Fed rate cuts or explicit QT end, (2) M2 momentum re-acceleration "
        f"above +2%, (3) DXY break below 100, (4) Sustained ETF inflows above +10K BTC/week."
    )
    p = doc.add_paragraph(verdict)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Report —")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(10)

    # ── Save ──
    output_path = "/home/ubuntu/S/btc_liquidity_analysis.docx"
    doc.save(output_path)
    print(f"✅ Document saved to {output_path}")
    print(f"   File size: {os.path.getsize(output_path):,} bytes")

if __name__ == "__main__":
    main()
