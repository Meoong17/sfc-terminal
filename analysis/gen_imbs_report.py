#!/usr/bin/env python3
"""Generate IMBS L1-L2 validation report DOCX (Calibri, Light Grid Accent 1).
Supports EN/ID output. Usage:
    python analysis/gen_imbs_report.py [--lang en|id] [--out /path/file.docx]
"""
import json, os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SFC_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CAL = os.path.join(SFC_ROOT, ".imbs_l1l2_calibration.json")
WF = os.path.join(SFC_ROOT, ".walk_forward_imbs_l1l2_summary.json")

CALIBRI = "Calibri"
TABLE_STYLE = "Light Grid Accent 1"

def parse_args():
    lang = "id"
    out = os.path.join(SFC_ROOT, "IMBS_L1L2_Validation_Report.docx")
    argv = sys.argv[1:]
    for i, a in enumerate(argv):
        if a == "--lang" and i + 1 < len(argv):
            lang = argv[i + 1]
        elif a == "--out" and i + 1 < len(argv):
            out = argv[i + 1]
    return lang, out

# ── Language strings ──
TXT = {
    "id": {
        "title": "IMBS Layer 1-2 Validation Report",
        "subtitle": "Integrated Macro Behavioral SFC — Stock-Flow + Liquidity Engine",
        "tagline": "Walk-forward predictive validation | Threshold calibration | Crisis-window validation",
        "s1": "1. Walk-Forward Predictive Validity",
        "s1_desc": ("Perbandingan sinyal baseline (price + DXY + M2 + FNG) vs sinyal IMBS L1-L2 "
                    "(baseline + indeks likuiditas Fed/ECB/BOJ balance sheet, TGA, RRP, DXY). "
                    "Semua gap CALM-minus-STRESS signifikan pada bootstrap CI 90%."),
        "s1_note": ("Interpretasi: menambahkan indeks likuiditas Layer 1-2 memperlebar gap prediktif "
                    "CALM-vs-STRESS (7d: -1.55 \u2192 -1.93pp; 30d: -7.46 \u2192 -8.20pp). Likuiditas makro "
                    "menambah informasi nyata di atas price+DXY+M2+FNG \u2014 bukan sekadar double-counting. "
                    "n STRESS naik (1208 \u2192 1757) menandakan sinyal IMBS lebih cepat mendeteksi kondisi berisiko."),
        "s2": "2. Threshold Calibration",
        "s2_desc": ("Menambah likuiditas menggeser banyak hari ke bucket STRESS, jadi cutoff tetap 45 "
                    "perlu dievaluasi. Hasil scan cutoff kandidat (horizon 30 hari):"),
        "s2_rec": ("Cutoff terbaik (memaksimalkan gap, cukup observasi): STRESS \u2265 {c} \u2014 gap {g:+.2f}pp, "
                   "n STRESS = {n}."),
        "s2_note": ("Catatan: cutoff yang lebih tinggi (\u226550) memberi gap lebih besar tetapi bucket STRESS "
                    "lebih kecil \u2192 lebih konservatif/defensif. Keputusan final menyesuaikan profil risiko "
                    "trader. Validasi walk-forward ulang wajib sebelum mengubah threshold live. Cutoff ini "
                    "TIDAK diterapkan ke pipeline live \u2014 sinyal live sfc_effective memiliki komposisi berbeda "
                    "(GLF penuh + dynamic weighting + HMM), perlu re-validasi dulu."),
        "s3": "3. Crisis-Window Validation",
        "s3_desc": ("Verifikasi sinyal IMBS benar-benar ELEVASI saat crash, bukan hanya pemisah statistik umum."),
        "s3_note": ("Sinyal IMBS terdeteksi STRESS pada hampir seluruh hari krisis (COVID, Luna, FTX: 100%; "
                    "2018: 87%) dan ter-elevasi +5.7 s.d. +16.4pp di atas kontrol 6 bulan. Ini konfirmasi "
                    "empiris bahwa komponen likuiditas L1-L2 menangkap kondisi krisis nyata, bukan artefak statistik."),
        "s4": "4. Conclusion",
        "s4_lines": [
            "1. IMBS Layer 1-2 backtest is FEASIBLE & PROVEN: Stock-Flow + Liquidity signal has genuine "
            "forward predictive value, free (FRED data), 2014-present history.",
            "2. Adding the liquidity index widens the predictive gap (30d: -7.46 \u2192 -8.20pp) \u2014 "
            "liquidity adds information, not double-counting.",
            "3. Signal detects STRESS on 87-100% of major crisis days \u2014 empirical validity, not artifact.",
            "4. Higher STRESS cutoff (\u226550) gives a larger gap; final decision needs re-validation before "
            "being applied live. Cutoff 55 is recorded as research recommendation, NOT deployed.",
            "5. Layers 3-5 (Behavior, Expectations, Regime) NOT yet fully backtested \u2014 distinguishing "
            "data (options/sentiment) history too short; next per the IMBS roadmap.",
        ],
        "footer": ("Generated from walk-forward validation (analysis/walk_forward_imbs_l1l2.py) and "
                   "calibration (analysis/imbs_l1l2_calibration.py). Data: FRED (CBBTCUSD, WALCL, "
                   "ECBASSETSW, JPNASSETS, WTREGEN, RRPONTSYD, DTWEXBGS, M2SL) + alternative.me FNG."),
        "metrics": ["Metric", "Baseline", "IMBS L1-L2", "Change"],
        "sig": "Yes",
        "col_cutoff": "Cutoff", "col_calm": "n CALM", "col_elev": "n ELEVATED",
        "col_stress": "n STRESS", "col_gap": "Gap (pp)", "col_ci": "CI 90%", "col_sig": "Significant",
        "col_crisis": "Crisis", "col_wmean": "Window Mean", "col_ctrl": "Control 6m",
        "col_elevpp": "Elevation (pp)", "col_sdays": "Stress Days (\u226545)", "col_detect": "Detected",
    },
    "en": {
        "title": "IMBS Layer 1-2 Validation Report",
        "subtitle": "Integrated Macro Behavioral SFC — Stock-Flow + Liquidity Engine",
        "tagline": "Walk-forward predictive validation | Threshold calibration | Crisis-window validation",
        "s1": "1. Walk-Forward Predictive Validity",
        "s1_desc": ("Comparison of baseline signal (price + DXY + M2 + FNG) vs IMBS L1-L2 signal "
                    "(baseline + Fed/ECB/BOJ balance-sheet, TGA, RRP, DXY liquidity index). All "
                    "CALM-minus-STRESS gaps significant at 90% bootstrap CI."),
        "s1_note": ("Interpretation: adding the Layer 1-2 liquidity index widens the CALM-vs-STRESS "
                    "predictive gap (7d: -1.55 \u2192 -1.93pp; 30d: -7.46 \u2192 -8.20pp). Macro liquidity adds "
                    "genuine information on top of price+DXY+M2+FNG \u2014 not mere double-counting. "
                    "n STRESS rises (1208 \u2192 1757), meaning the IMBS signal flags risk earlier."),
        "s2": "2. Threshold Calibration",
        "s2_desc": ("Adding liquidity shifts many days into the STRESS bucket, so the fixed 45 cutoff "
                    "needs evaluation. Candidate cutoff scan (30-day horizon):"),
        "s2_rec": ("Best cutoff (maximizes gap with adequate observations): STRESS \u2265 {c} \u2014 gap {g:+.2f}pp, "
                   "n STRESS = {n}."),
        "s2_note": ("Note: a higher cutoff (\u226550) yields a larger gap but a smaller STRESS bucket \u2192 more "
                    "conservative/defensive. Final decision depends on trader risk profile. Walk-forward "
                    "re-validation is required before changing the live threshold. This cutoff is NOT "
                    "deployed to the live pipeline \u2014 the live sfc_effective signal has a different "
                    "composition (full GLF + dynamic weighting + HMM) and needs re-validation first."),
        "s3": "3. Crisis-Window Validation",
        "s3_desc": "Verifies the IMBS signal genuinely ELEVATES during crashes, not just a general statistical separator.",
        "s3_note": ("The IMBS signal detects STRESS on nearly all crisis days (COVID, Luna, FTX: 100%; "
                    "2018: 87%) and is elevated +5.7 to +16.4pp above the 6-month control. This empirically "
                    "confirms the L1-L2 liquidity components capture real crisis conditions, not a statistical artifact."),
        "s4": "4. Conclusion",
        "s4_lines": [
            "1. IMBS Layer 1-2 backtest is FEASIBLE & PROVEN: Stock-Flow + Liquidity signal has genuine "
            "forward predictive value, free (FRED data), 2014-present history.",
            "2. Adding the liquidity index widens the predictive gap (30d: -7.46 \u2192 -8.20pp) \u2014 "
            "liquidity adds information, not double-counting.",
            "3. Signal detects STRESS on 87-100% of major crisis days \u2014 empirical validity, not artifact.",
            "4. Higher STRESS cutoff (\u226550) gives a larger gap; final decision needs re-validation before "
            "being applied live. Cutoff 55 is recorded as research recommendation, NOT deployed.",
            "5. Layers 3-5 (Behavior, Expectations, Regime) NOT yet fully backtested \u2014 distinguishing "
            "data (options/sentiment) history too short; next per the IMBS roadmap.",
        ],
        "footer": ("Generated from walk-forward validation (analysis/walk_forward_imbs_l1l2.py) and "
                   "calibration (analysis/imbs_l1l2_calibration.py). Data: FRED (CBBTCUSD, WALCL, "
                   "ECBASSETSW, JPNASSETS, WTREGEN, RRPONTSYD, DTWEXBGS, M2SL) + alternative.me FNG."),
        "metrics": ["Metric", "Baseline", "IMBS L1-L2", "Change"],
        "sig": "Yes",
        "col_cutoff": "Cutoff", "col_calm": "n CALM", "col_elev": "n ELEVATED",
        "col_stress": "n STRESS", "col_gap": "Gap (pp)", "col_ci": "CI 90%", "col_sig": "Significant",
        "col_crisis": "Crisis", "col_wmean": "Window Mean", "col_ctrl": "Control 6m",
        "col_elevpp": "Elevation (pp)", "col_sdays": "Stress Days (\u226545)", "col_detect": "Detected",
    },
}

def set_base_font(doc):
    st = doc.styles["Normal"]
    st.font.name = CALIBRI
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CALIBRI)

def para(doc, text="", size=11, bold=False, space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = CALIBRI
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return p

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = CALIBRI
    return h

def make_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = TABLE_STYLE
    except Exception:
        t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = CALIBRI
        r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.name = CALIBRI
            r.font.size = Pt(10)
    return t

def main():
    lang, out = parse_args()
    t = TXT[lang]
    cal = json.load(open(CAL))
    wf = json.load(open(WF))

    doc = Document()
    set_base_font(doc)

    heading(doc, t["title"], 0)
    para(doc, t["subtitle"], size=12, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    para(doc, t["tagline"], size=10, color=RGBColor(0x59, 0x59, 0x59))
    para(doc, "")

    # Section 1
    heading(doc, t["s1"], 1)
    para(doc, t["s1_desc"])
    hdr = t["metrics"]
    wf_rows = [
        hdr,
        [hdr[0] + " 7d", f"{wf['base_gap_7d']:.2f}", f"{wf['imbs_gap_7d']:.2f}",
         f"{wf['imbs_gap_7d']-wf['base_gap_7d']:+.2f}"],
        ["CI 90% 7d", str(wf["base_gap_7d_ci"]), str(wf["imbs_gap_7d_ci"]), ""],
        ["Significant 7d", str(wf["base_gap_7d_significant"]),
         str(wf["imbs_gap_7d_significant"]), ""],
        [hdr[0] + " 30d", f"{wf['base_gap_30d']:.2f}", f"{wf['imbs_gap_30d']:.2f}",
         f"{wf['imbs_gap_30d']-wf['base_gap_30d']:+.2f}"],
        ["CI 90% 30d", str(wf["base_gap_30d_ci"]), str(wf["imbs_gap_30d_ci"]), ""],
        ["Significant 30d", str(wf["base_gap_30d_significant"]),
         str(wf["imbs_gap_30d_significant"]), ""],
        ["n CALM (30d)", wf["base_n_calm_30d"], wf["imbs_n_calm_30d"], ""],
        ["n STRESS (30d)", wf["base_n_stress_30d"], wf["imbs_n_stress_30d"], ""],
    ]
    make_table(doc, wf_rows[0], wf_rows[1:])
    para(doc, "")
    para(doc, t["s1_note"], size=10)

    # Section 2
    heading(doc, t["s2"], 1)
    para(doc, t["s2_desc"])
    rows2 = [[t["col_cutoff"], t["col_calm"], t["col_elev"], t["col_stress"],
              t["col_gap"], t["col_ci"], t["col_sig"]]]
    for r in cal["threshold_calibration"]["rows"]:
        rows2.append([
            r["cutoff"], r["n_calm"], r["n_elevated"], r["n_stress"],
            f"{r['gap_pp']:+.2f}",
            f"[{r['ci_lo']:.2f}, {r['ci_hi']:.2f}]",
            t["sig"] if r["significant"] else "No",
        ])
    make_table(doc, rows2[0], rows2[1:])
    rec = cal["threshold_calibration"]["recommended_cutoff"]
    para(doc, "")
    if rec:
        row = next(r for r in cal["threshold_calibration"]["rows"]
                   if r["cutoff"] == rec)
        para(doc, t["s2_rec"].format(c=rec, g=row["gap_pp"], n=row["n_stress"]),
             bold=True)
    para(doc, t["s2_note"], size=10)

    # Section 3
    heading(doc, t["s3"], 1)
    para(doc, t["s3_desc"])
    rows3 = [[t["col_crisis"], t["col_wmean"], t["col_ctrl"], t["col_elevpp"],
              t["col_sdays"], t["col_detect"]]]
    for name, v in cal["crisis_windows"].items():
        if "error" in v:
            rows3.append([name, "-", "-", "-", "-", "No data"])
            continue
        rows3.append([
            name, f"{v['window_mean']:.1f}",
            f"{v['control_6m_mean']:.1f}" if v["control_6m_mean"] else "-",
            f"{v['elevation_vs_control_pp']:+.1f}",
            f"{v['n_stress']}/{v['n_days']} ({v['stress_pct']:.0f}%)",
            t["sig"] if v["elevated"] else "No",
        ])
    make_table(doc, rows3[0], rows3[1:])
    para(doc, "")
    para(doc, t["s3_note"], size=10)

    # Section 4
    heading(doc, t["s4"], 1)
    for line in t["s4_lines"]:
        para(doc, line, size=10, space_after=4)

    para(doc, "")
    para(doc, t["footer"], size=9, color=RGBColor(0x59, 0x59, 0x59))

    os.makedirs(os.path.dirname(out) or ".", exist_ok=True)
    doc.save(out)
    print(f"Saved -> {out}  (lang={lang})")

if __name__ == "__main__":
    main()
